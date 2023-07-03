import openai
import difflib
from docx import Document
from docx.shared import RGBColor

def create_doc_with_diff(diff_data, filename):
    doc = Document()

    for original_chunk, updated_chunk, diff in diff_data:
        doc.add_paragraph('Original Chunk:')
        doc.add_paragraph(original_chunk)

        doc.add_paragraph('Updated Chunk:')
        doc.add_paragraph(updated_chunk)

        doc.add_paragraph('Differences:')
        for line in diff:
            if line.startswith('- '):  
                run = doc.add_paragraph().add_run(line)
                run.underline = True
                run.font.color.rgb = RGBColor(255, 0, 0)  # red color for deleted lines
            elif line.startswith('+ '):  
                run = doc.add_paragraph().add_run(line)
                run.underline = True
                run.font.color.rgb = RGBColor(0, 255, 0)  # green color for added lines
            else:
                doc.add_paragraph(line)
        doc.add_paragraph('-----------------------------------')

    doc.save(filename)


def process_chunks(original_chunks, update_chunks):
    openai.api_key = 'YOUR_OPENAI_API_KEY'
    diff_data = []

    for original_chunk, update_chunk in zip(original_chunks, update_chunks):
        response = openai.Completion.create(
            model="gpt-3.5-turbo",
            messages=[
                {
                    "role": "system",
                    "content": "You are an AI trained to process and update regulation data chunks. If an incoming update data chunk corresponds appropriately to any segment within the original regulation data chunk, perform the update. If the update data chunk doesn't align or isn't compatible, keep the original regulation data chunk as is."
                },
                {
                    "role": "user",
                    "content": f"Here is the original regulation data chunk: {original_chunk}"
                },
                {
                    "role": "user",
                    "content": f"Here is the update data chunk: {update_chunk}"
                },
                {
                    "role": "user",
                    "content": "Please compare the update data chunk with the original regulation data chunk and perform an update if appropriate. Return the resultant data chunk."
                }
            ],
            options={
                "use_cache": False,
                "max_tokens": 4000
            }
        )

        updated_text = response['choices'][0]['message']['content']
        if original_chunk != updated_text:
            diff = list(difflib.ndiff(original_chunk.splitlines(), updated_text.splitlines()))
            diff_data.append((original_chunk, updated_text, diff))

    return diff_data


def main():
    original_chunks = ["Your original 4000 tokens data chunk1", "Your original 4000 tokens data chunk2"]  # List of your original chunks
    update_chunks = ["Your 4000 tokens update data chunk1", "Your 4000 tokens update data chunk2"]  # List of your update chunks

    diff_data = process_chunks(original_chunks, update_chunks)
    if diff_data:
        create_doc_with_diff(diff_data, "updated_chunks.docx")


if __name__ == "__main__":
    main()
