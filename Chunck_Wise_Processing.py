import openai
import difflib
from docx import Document
from docx.shared import RGBColor

def create_doc_with_diff(original_text, updated_text, filename):
    diff = difflib.ndiff(original_text.splitlines(), updated_text.splitlines())
    doc = Document()

    for i, line in enumerate(diff):
        if line.startswith('- '):  
            run = doc.add_paragraph().add_run(line)
            run.underline = True
            run.font.color.rgb = RGBColor(255, 0, 0)  # red color for deleted lines
        elif line.startswith('+ '):  
            run = doc.add_paragraph().add_run(line)
            run.underline = True
            run.font.color.rgb = RGBColor(0, 255, 0)  # green color for added lines
        else:
            doc.add_paragraph(line)

    doc.save(filename)


def process_chunks(original_chunk, update_chunk):
    openai.api_key = 'YOUR_OPENAI_API_KEY'

    response = openai.Completion.create(
        model="gpt-3.5-turbo",
        messages=[
            {
                "role": "system",
                "content": "You are an AI trained to process and update regulation data chunks. If an incoming update data chunk corresponds appropriately to any segment within the original regulation data chunk, perform the update. If the update data chunk doesn't align or isn't compatible, keep the original regulation data chunk as is."
            },
            {
                "role": "user",
                "content": f"Here is the original regulation data chunk: {original_chunk}"
            },
            {
                "role": "user",
                "content": f"Here is the update data chunk: {update_chunk}"
            },
            {
                "role": "user",
                "content": "Please compare the update data chunk with the original regulation data chunk and perform an update if appropriate. Return the resultant data chunk."
            }
        ],
        options={
            "use_cache": False,
            "max_tokens": 4000
        }
    )

    return response['choices'][0]['message']['content']


def main():
    original_chunk = "Your original 4000 tokens data chunk"
    update_chunk = "Your 4000 tokens update data chunk"

    updated_text = process_chunks(original_chunk, update_chunk)

    if original_chunk != updated_text:
        create_doc_with_diff(original_chunk, updated_text, "updated_chunk.docx")


if __name__ == "__main__":
    main()
